# utils/document_loader.py
import os
import fitz  # PyMuPDF
import re
from langchain.docstore.document import Document
from utils.category_manager import assign_category
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# -------------------------------
# Text in strukturierte Abschnitte (Chunks) teilen
# -------------------------------
def split_into_chunks_by_heading(text):
    pattern = r"(?=^\s*(§{1,2}\s*\d+[^\n]*|Art\.?\s*\d+[^\n]*|Ziff\.?\s*\d+[^\n]*|Artikel\s+\d+[^\n]*))"
    chunks = re.split(pattern, text, flags=re.MULTILINE)

    grouped = []
    for i in range(0, len(chunks) - 1, 2):
        heading = chunks[i].strip()
        content = chunks[i + 1].strip()
        grouped.append({"heading": heading, "content": content})

    if len(chunks) % 2 == 1:
        grouped.append({"heading": "–", "content": chunks[-1].strip()})
    return grouped

# -------------------------------
# PDF-Dateien laden & chunken
# -------------------------------
def extract_chunks_from_pdf(path):
    try:
        doc = fitz.open(path)
    except Exception as e:
        print(f"Fehler beim Öffnen von PDF {path}: {e}")
        return []

    all_chunks = []
    for i, page in enumerate(doc, start=1):
        text = page.get_text()
        page_chunks = split_into_chunks_by_heading(text)
        for chunk in page_chunks:
            all_chunks.append(Document(
                page_content=chunk["content"],
                metadata={
                    "source": os.path.basename(path),
                    "page": i,
                    "heading": chunk["heading"]
                }
            ))
    return all_chunks

# -------------------------------
# DOCX-Dateien laden & chunken
# -------------------------------
def extract_chunks_from_docx(path):
    try:
        doc = DocxDocument(path)
    except Exception as e:
        print(f"Fehler beim Laden von DOCX {path}: {e}")
        return []

    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    chunks = split_into_chunks_by_heading(full_text)
    all_chunks = []

    for chunk in chunks:
        all_chunks.append(Document(
            page_content=chunk["content"],
            metadata={
                "source": os.path.basename(path),
                "page": None,
                "heading": chunk["heading"]
            }
        ))

    # Vorschau erzeugen
    export_docx_to_pdf(path)
    return all_chunks

# -------------------------------
# Lade alle Dokumente im Ordner & Kategorie zuweisen
# -------------------------------
def load_documents_from_folder(folder_path):
    documents = []
    if not os.path.exists(folder_path):
        print(f"Ordner {folder_path} existiert nicht.")
        return []

    for filename in os.listdir(folder_path):
        full_path = os.path.join(folder_path, filename)
        ext = os.path.splitext(filename)[-1].lower()

        chunks = []
        if ext == ".pdf":
            chunks = extract_chunks_from_pdf(full_path)
        elif ext == ".docx":
            chunks = extract_chunks_from_docx(full_path)
        else:
            continue

        for chunk in chunks:
            category = assign_category(chunk.page_content)
            chunk.metadata["category"] = category
            documents.append(chunk)

    return documents

# -------------------------------
# DOCX zu PDF Vorschau exportieren
# -------------------------------
def export_docx_to_pdf(docx_path, output_dir="previews"):
    os.makedirs(output_dir, exist_ok=True)
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        print(f"Fehler beim Export von DOCX {docx_path}: {e}")
        return None

    pdf_path = os.path.join(output_dir, os.path.basename(docx_path).replace(".docx", ".pdf"))
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 40

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y -= 15
            continue
        for line in text.split("\n"):
            c.drawString(40, y, line[:120])
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 40

    c.save()
    return pdf_path
